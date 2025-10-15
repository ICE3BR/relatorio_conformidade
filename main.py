import pandas as pd
from docx import Document
from datetime import datetime
import os
import re


def preencher_relatorio():
    # Definir caminhos dos arquivos
    diretorio_atual = os.path.dirname(os.path.abspath(__file__))
    arquivo_excel = os.path.join(diretorio_atual, 'Conformidade - Legal OpsTESTE.xlsx')
    arquivo_word = os.path.join(diretorio_atual, 'MODELO RELATORIO.docx')
    
    # Verificar se os arquivos existem
    if not os.path.exists(arquivo_excel):
        print(f"ERRO: Arquivo Excel não encontrado: {arquivo_excel}")
        print("Certifique-se de que o arquivo 'Conformidade - Legal OpsTESTE.xlsx' está na mesma pasta do script.")
        return
    
    if not os.path.exists(arquivo_word):
        print(f"ERRO: Arquivo Word não encontrado: {arquivo_word}")
        print("Certifique-se de que o arquivo 'TESTE RELATORIO.docx' está na mesma pasta do script.")
        return
    
    # Ler a planilha Excel
    try:
        df = pd.read_excel(arquivo_excel)
        print(f"Planilha carregada com {len(df)} processos")
    except Exception as e:
        print(f"Erro ao ler a planilha Excel: {e}")
        return
    
    # Mapeamento de colunas para placeholders {CHAVE} (substituição direta)
    mapeamento = {
        'NUMERO_PROCESSO': 'NUMERO_PROCESSO',
        'AUTOR': 'AUTOR',
        'CUMPRIMENTO_SENTENCA': 'CUMPRIMENTO_SENTENCA',
        'SITUACAO_PROCESSO': 'SITUACAO_PROCESSO',
        'DATA_ACAO': 'DATA_ACAO',
        'DATA_PERICIA': 'DATA_PERICIA',
        'DATA_REALIZADA': 'DATA_REALIZADA',
        'DATA_LAUDO': 'DATA_LAUDO',
        'TIPO LAUDO': 'TIPO LAUDO',
        'DATA_SENTENCA': 'DATA_SENTENCA',
        'SENTENCA': 'SENTENCA',
        'DATA_APELACAO': 'DATA_APELACAO',
        'APE': 'APE',
        'DATA_JULGAMENTO': 'DATA_JULGAMENTO',
        'JULGA': 'JULGA',
        'DATA_TRANSITO': 'DATA_TRANSITO',
        'DATA_CUMPRIMENTO': 'DATA_CUMPRIMENTO',
        'DATA_HOMOLOGACAO': 'DATA_HOMOLOGACAO',
        'DATA_PRECA': 'DATA_PRECA',
        'DATA_OFICIO': 'DATA_OFICIO',
        'DATA_OR_PAGAMENTO': 'DATA_OR_PAGAMENTO',
        'DATA_ENCERRAMENTO': 'DATA_ENCERRAMENTO'
    }

    # --- Helpers ---
    def formatar_data(valor):
        """Formata datas para o padrão DD/MM/AAAA"""
        if pd.isna(valor) or valor == '' or valor == 'None' or valor is None:
            return ''
        try:
            if isinstance(valor, str):
                valor = valor.strip()
                # tenta 'YYYY-MM-DD HH:MM:SS' ou 'YYYY-MM-DD'
                if re.match(r'^\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}$', valor):
                    data_obj = datetime.strptime(valor, '%Y-%m-%d %H:%M:%S')
                elif re.match(r'^\d{4}-\d{2}-\d{2}$', valor):
                    data_obj = datetime.strptime(valor, '%Y-%m-%d')
                else:
                    # tenta interpretar outros formatos
                    data_obj = pd.to_datetime(valor, dayfirst=False, errors='raise').to_pydatetime()
            else:
                # pandas Timestamp / datetime
                data_obj = pd.to_datetime(valor, errors='raise').to_pydatetime()
            return data_obj.strftime('%d/%m/%Y')
        except Exception:
            return str(valor)

    def limpar_nome_arquivo(nome):
        """Remove caracteres inválidos para nome de arquivo"""
        caracteres_invalidos = ['<', '>', ':', '"', '/', '\\', '|', '?', '*']
        for char in caracteres_invalidos:
            nome = nome.replace(char, '-')
        return nome

    def norm(txt: str) -> str:
        """Normaliza para comparação (lower + remove acentos, espaços extras)."""
        if txt is None or (isinstance(txt, float) and pd.isna(txt)):
            return ''
        s = str(txt).strip().lower()
        # remoção simples de acentos
        s = (s.replace('á', 'a').replace('à', 'a').replace('â', 'a').replace('ã', 'a')
               .replace('é', 'e').replace('ê', 'e')
               .replace('í', 'i')
               .replace('ó', 'o').replace('ô', 'o').replace('õ', 'o')
               .replace('ú', 'u')
               .replace('ç', 'c'))
        return s

    def aplicar_marcacoes(texto: str, linha: pd.Series) -> str:
        """
        Marca X nos placeholders especiais, substituindo cada token por (X) ou ( ).
        Tokens esperados no .docx: ({LP}) ({LPP}) ({LN}) ({SENTENCA_A}) ({SENTENCA_I})
                                ({APE_A}) ({APE_I}) ({JULGA_A}) ({JULGA_I})
        Regras:
        - Laudo: usa 'TIPO DE LAUDO' (fallback 'LAUDO')
            * positivo  -> LP
            * parcial   -> LPP
            * negativo  -> LN
        - Sentença: 'procedente' -> SENTENCA_A ; 'improcedente' -> SENTENCA_I
        - Apelação: 'autor' -> APE_A ; 'inss' -> APE_I
        - Julgamento: 'favoravel' -> JULGA_A ; 'desfavoravel' -> JULGA_I
        """
        if not texto:
            return texto

        # Começa assumindo que todos estão desmarcados
        marcacoes = {
            'LP': '( )',
            'LPP': '( )',
            'LN': '( )',
            'SENTENCA_A': '( )',
            'SENTENCA_I': '( )',
            'APE_A': '( )',
            'APE_I': '( )',
            'JULGA_A': '( )',
            'JULGA_I': '( )',
        }

        # --- Laudo ---
        laudo_bruto = linha.get('TIPO LAUDO', '')
        if pd.isna(laudo_bruto) or laudo_bruto == '':
            laudo_bruto = linha.get('LAUDO', '')
        
        laudo_n = norm(laudo_bruto)
        
        # Verifica cada tipo de laudo
        if 'laudo positivo -  1º grau' in laudo_n:
            marcacoes['LP'] = '(X)'
        elif 'laudo parcial -  1º grau' in laudo_n:
            marcacoes['LPP'] = '(X)'
        elif 'laudo negativo  - 1º grau' in laudo_n:
            marcacoes['LN'] = '(X)'

        # --- Sentença ---
        sentenca_bruto = linha.get('SENTENÇA', '')
        if pd.isna(sentenca_bruto) or sentenca_bruto == '':
            sentenca_bruto = linha.get('SENTENCA', '')
        
        sentenca_n = norm(sentenca_bruto)
        
        if 'procedente' in sentenca_n:
            marcacoes['SENTENCA_A'] = '(X)'
        elif 'improcedent e' in sentenca_n:
            marcacoes['SENTENCA_I'] = '(X)'

        # --- Apelação ---
        apelacao_bruto = linha.get('APELAÇÃO', '')
        if pd.isna(apelacao_bruto) or apelacao_bruto == '':
            apelacao_bruto = linha.get('APELACAO', linha.get('APE', ''))
        
        apelacao_n = norm(apelacao_bruto)
        
        if 'autor' in apelacao_n:
            marcacoes['APE_A'] = '(X)'
        elif '' in apelacao_n:
            marcacoes['APE_I'] = '(X)'

        # --- Julgamento ---
        julgamento_bruto = linha.get('JULGAMENTO', '')
        if pd.isna(julgamento_bruto) or julgamento_bruto == '':
            julgamento_bruto = linha.get('JULGA', '')
        
        julgamento_n = norm(julgamento_bruto)
        
        if 'favoravel' in julgamento_n:
            marcacoes['JULGA_A'] = '(X)'
        elif 'desfavoravel' in julgamento_n:
            marcacoes['JULGA_I'] = '(X)'

        # Troca todos os tokens no texto pelo respectivo (X) ou ( )
        for token, repl in marcacoes.items():
            texto = texto.replace(f'({{{token}}})', repl)

        return texto

    # Para cada processo na planilha
    for index, row in df.iterrows():
        numero_processo = str(row['NUMERO_PROCESSO']) if 'NUMERO_PROCESSO' in row else f'_{index+1:03d}'
        print(f"Processando processo {index + 1}/{len(df)}: {numero_processo}")
        
        try:
            # Carregar o template Word
            doc = Document(arquivo_word)
            
            # Substituir nos parágrafos
            for paragraph in doc.paragraphs:
                texto_original = paragraph.text
                texto_substituido = texto_original
                
                # Substituições diretas {CHAVE}
                for coluna, placeholder in mapeamento.items():
                    chave = f'{{{placeholder}}}'
                    if chave in texto_substituido and coluna in row:
                        valor = row[coluna]
                        if 'DATA' in coluna:
                            valor = formatar_data(valor)
                        else:
                            valor = '' if pd.isna(valor) or valor in ('', 'None', None) else str(valor)
                        texto_substituido = texto_substituido.replace(chave, valor)

                # Marcações especiais -> coloca X dentro dos parênteses quando bater
                novo_texto = aplicar_marcacoes(texto_substituido, row)

                if novo_texto != texto_original:
                    # Limpa os runs existentes e insere o novo texto
                    for r in paragraph.runs:
                        r.text = ''
                    paragraph.add_run(novo_texto)
            
            # Substituir nas tabelas
            for table in doc.tables:
                for row_table in table.rows:
                    for cell in row_table.cells:
                        texto_original = cell.text
                        texto_substituido = texto_original
                        
                        # Substituições diretas {CHAVE}
                        for coluna, placeholder in mapeamento.items():
                            chave = f'{{{placeholder}}}'
                            if chave in texto_substituido and coluna in row:
                                valor = row[coluna]
                                if 'DATA' in coluna:
                                    valor = formatar_data(valor)
                                else:
                                    valor = '' if pd.isna(valor) or valor in ('', 'None', None) else str(valor)
                                texto_substituido = texto_substituido.replace(chave, valor)

                        # Marcações especiais -> coloca X dentro dos parênteses quando bater
                        texto_substituido = aplicar_marcacoes(texto_substituido, row)
                        
                        if texto_substituido != texto_original:
                            cell.text = texto_substituido
            
            # Salvar o documento preenchido
            nome_arquivo_saida = f'RELATORIO_{limpar_nome_arquivo(numero_processo)}.docx'
            caminho_saida = os.path.join(diretorio_atual, nome_arquivo_saida)
            doc.save(caminho_saida)
            print(f'✓ Relatório gerado: {nome_arquivo_saida}')
            
        except Exception as e:
            print(f"✗ Erro ao processar processo {numero_processo}: {e}")
    
    print("\nProcessamento concluído!")


if __name__ == "__main__":
    preencher_relatorio()
