import pandas as pd 
from docx import Document 
from docx.shared import Pt 
from docx.oxml.ns import qn 
from datetime import datetime, timedelta
import os 
import re 
from dateutil.relativedelta import relativedelta

def preencher_relatorio(): 
    # Definir caminhos dos arquivos 
    diretorio_atual = os.path.dirname(os.path.abspath(__file__)) 
    arquivo_excel = os.path.join(diretorio_atual, 'Pasta1 teste.xlsx') 
    arquivo_word = os.path.join(diretorio_atual, 'MODELO RELATORIO.docx') 

    # Verificar se os arquivos existem 
    if not os.path.exists(arquivo_excel): 
        print(f"ERRO: Arquivo Excel não encontrado: {arquivo_excel}") 
        print("Certifique-se de que o arquivo 'Conformidade - xlsx' está na mesma pasta do script.") 
        return 

    if not os.path.exists(arquivo_word): 
        print(f"ERRO: Arquivo Word não encontrado: {arquivo_word}") 
        print("Certifique-se de que o arquivo 'TESTE RELATORIO.docx' está na mesma pasta do script.") 
        return 

    # Ler a planilha Excel 
    try: 
        df = pd.read_excel(arquivo_excel) 
        print(f"Planilha carregada com {len(df)} processos") 
    except Exception as e: 
        print(f"Erro ao ler a planilha Excel: {e}") 
        return 

    # Mapeamento de colunas para placeholders {CHAVE} (substituição direta) 
    mapeamento = { 
        'NUMERO_PROCESSO': 'NUMERO_PROCESSO', 
        'AUTOR': 'AUTOR', 
        'CUMPRIMENTO_SENTENCA': 'CUMPRIMENTO_SENTENCA', 
        'SITUACAO_PROCESSO': 'SITUACAO_PROCESSO', 
        'DATA_ACAO': 'DATA_ACAO', 
        'DATA_PERICIA': 'DATA_PERICIA', 
        'DATA_REALIZADA': 'DATA_REALIZADA', 
        'DATA_LAUDO': 'DATA_LAUDO', 
        'TIPO LAUDO': 'TIPO LAUDO', 
        'DATA_SENTENCA': 'DATA_SENTENCA', 
        'SENTENCA': 'SENTENCA', 
        'DATA_APELACAO': 'DATA_APELACAO', 
        'APE': 'APE', 
        'DATA_JULGAMENTO': 'DATA_JULGAMENTO', 
        'JULGA': 'JULGA', 
        'DATA_TRANSITO': 'DATA_TRANSITO', 
        'DATA_CUMPRIMENTO': 'DATA_CUMPRIMENTO', 
        'DATA_HOMOLOGACAO': 'DATA_HOMOLOGACAO', 
        'DATA_PRECA': 'DATA_PRECA', 
        'DATA_OFICIO': 'DATA_OFICIO', 
        'DATA_OR_PAGAMENTO': 'DATA_OR_PAGAMENTO', 
        'DATA_ENCERRAMENTO': 'DATA_ENCERRAMENTO' 
    } 

    # Sequência de eventos com prazos em anos meses e dias
    EVENTOS_SEQUENCIA = [
        ("Distribuição da ação", 0, 0, 0),
        ("Designação de perícia médica judicial", 0, 2, 0),
        ("Realização da perícia médica", 0, 1, 20),
        ("Laudo pericial", 0, 2, 0),
        ("Sentença Proferida", 0, 3, 0),
        ("Apelação interposta", 0, 1, 15),
        ("Julgamento", 0, 2, 0),
        ("Trânsito em julgado", 0, 1, 0),
        ("Cumprimento de Sentença", 0, 1, 1),
        ("Homologação dos Valores", 0, 3, 0),
        ("Instauração do RPV ou Precatório", 0, 1, 5),
        ("Expedição de Ofício", 0, 3, 0),
        ("Ordem de Pagamento", 0, 1, 0),
        ("Encerramento com a Liberação do Valores", 1, 6, 0),
    ]

    # --- Helpers --- 
    def calcular_data_prevista(data_inicio, anos, meses, dias):
        """Calcula data prevista somando meses e dias à data de início"""
        if pd.isna(data_inicio) or data_inicio == '' or data_inicio == 'None' or data_inicio is None:
            return ''
        try:
            if isinstance(data_inicio, str):
                # Tenta converter string para datetime
                data_inicio = pd.to_datetime(data_inicio, errors='raise').to_pydatetime()
            
            # Calcula data prevista
            data_prevista = data_inicio + relativedelta(years=+anos, months=+meses, days=+dias)
            return data_prevista.strftime('%d/%m/%Y')
        except Exception:
            return ''

    def calcular_datas_previstas_sequencia(data_acao):
        """Calcula todas as datas previstas da sequência baseado na data da ação"""
        datas_previstas = {}
        
        if pd.isna(data_acao) or data_acao == '' or data_acao == 'None' or data_acao is None:
            return datas_previstas
            
        try:
            if isinstance(data_acao, str):
                data_atual = pd.to_datetime(data_acao, errors='raise').to_pydatetime()
            else:
                data_atual = pd.to_datetime(data_acao).to_pydatetime()
            
            for evento, anos, meses, dias in EVENTOS_SEQUENCIA:
                data_prevista = data_atual + relativedelta(years=+ anos, months=+meses, days=+dias)
                datas_previstas[evento] = data_prevista.strftime('%d/%m/%Y')
                data_atual = data_prevista
                
        except Exception as e:
            print(f"Erro ao calcular datas previstas: {e}")
            
        return datas_previstas

    def formatar_data(valor): 
        """Formata datas para o padrão DD/MM/AAAA""" 
        if pd.isna(valor) or valor == '' or valor == 'None' or valor is None: 
            return '' 
        try: 
            if isinstance(valor, str): 
                valor = valor.strip() 
                # tenta 'YYYY-MM-DD HH:MM:SS' ou 'YYYY-MM-DD' 
                if re.match(r'^\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}:\d{2}$', valor): 
                    data_obj = datetime.strptime(valor, '%Y-%m-%d %H:%M:%S') 
                elif re.match(r'^\d{4}-\d{2}-\d{2}$', valor): 
                    data_obj = datetime.strptime(valor, '%Y-%m-%d') 
                else: 
                    # tenta interpretar outros formatos 
                    data_obj = pd.to_datetime(valor, dayfirst=False, errors='raise').to_pydatetime() 
            else: 
                # pandas Timestamp / datetime 
                data_obj = pd.to_datetime(valor, errors='raise').to_pydatetime() 
            return data_obj.strftime('%d/%m/%Y') 
        except Exception: 
            return str(valor) 

    def limpar_nome_arquivo(nome): 
        """Remove caracteres inválidos para nome de arquivo""" 
        caracteres_invalidos = ['<', '>', ':', '"', '/', '\\', '|', '?', '*'] 
        for char in caracteres_invalidos: 
            nome = nome.replace(char, '-') 
        return nome 

    def norm(txt: str) -> str: 
        """Normaliza para comparação (lower + remove acentos, espaços extras).""" 
        if txt is None or (isinstance(txt, float) and pd.isna(txt)): 
            return '' 
        s = str(txt).strip().lower() 
        # remoção simples de acentos 
        s = (s.replace('á', 'a').replace('à', 'a').replace('â', 'a').replace('ã', 'a') 
             .replace('é', 'e').replace('ê', 'e') 
             .replace('í', 'i') 
             .replace('ó', 'o').replace('ô', 'o').replace('õ', 'o') 
             .replace('ú', 'u') 
             .replace('ç', 'c')) 
        return s 

    def aplicar_marcacoes(texto: str, linha: pd.Series) -> str: 
        """Marcação X nos placeholders especiais"""
        if not texto: 
            return texto 
        
        marcacoes = { 
            'LP': '( )', 
            'LPP': '( )', 
            'LN': '( )', 
            'SENTENCA_A': '( )', 
            'SENTENCA_I': '( )', 
            'APE_A': '( )', 
            'APE_I': '( )', 
            'JULGA_A': '( )', 
            'JULGA_I': '( )', 
        } 

        # --- Laudo --- 
        laudo_bruto = linha.get('TIPO LAUDO', '') 
        if pd.isna(laudo_bruto) or laudo_bruto == '': 
            laudo_bruto = linha.get('LAUDO', '') 
        laudo_n = norm(laudo_bruto) 
        
        if 'laudo positivo - 1º grau' in laudo_n: 
            marcacoes['LP'] = '(X)' 
        elif 'laudo parcial - 1º grau' in laudo_n: 
            marcacoes['LPP'] = '(X)' 
        elif 'laudo negativo - 1º grau' in laudo_n: 
            marcacoes['LN'] = '(X)' 

        # --- Sentença --- 
        sentenca_bruto = linha.get('SENTENÇA', '') 
        if pd.isna(sentenca_bruto) or sentenca_bruto == '': 
            sentenca_bruto = linha.get('SENTENCA', '') 
        sentenca_n = norm(sentenca_bruto) 
        
        if 'procedente' in sentenca_n: 
            marcacoes['SENTENCA_A'] = '(X)' 
        elif 'improcedent e' in sentenca_n: 
            marcacoes['SENTENCA_I'] = '(X)' 

        # --- Apelação --- 
        apelacao_bruto = linha.get('APELAÇÃO', '') 
        if pd.isna(apelacao_bruto) or apelacao_bruto == '': 
            apelacao_bruto = linha.get('APELACAO', linha.get('APE', '')) 
        apelacao_n = norm(apelacao_bruto) 
        
        if 'autor' in apelacao_n: 
            marcacoes['APE_A'] = '(X)' 
        elif '' in apelacao_n: 
            marcacoes['APE_I'] = '(X)' 

        # --- Julgamento --- 
        julgamento_bruto = linha.get('JULGAMENTO', '') 
        if pd.isna(julgamento_bruto) or julgamento_bruto == '': 
            julgamento_bruto = linha.get('JULGA', '') 
        julgamento_n = norm(julgamento_bruto) 
        
        if 'favoravel' in julgamento_n: 
            marcacoes['JULGA_A'] = '(X)' 
        elif 'desfavoravel' in julgamento_n: 
            marcacoes['JULGA_I'] = '(X)' 

        # Troca todos os tokens no texto pelo respectivo (X) ou ( ) 
        for token, repl in marcacoes.items(): 
            texto = texto.replace(f'({{{token}}})', repl) 
            
        return texto 

    def aplicar_fonte_calibri_light(run, cor_vermelha=False): 
        """Aplica fonte Calibri Light 10.5 ao run, opcionalmente em vermelho""" 
        run.font.name = 'Calibri Light' 
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri Light') 
        run.font.size = Pt(10.5)
        if cor_vermelha:
            from docx.shared import RGBColor
            run.font.color.rgb = RGBColor(255, 0, 0)  # Vermelho

    # --- Mapeamento de eventos para colunas ---
    mapeamento_eventos_colunas = {
        "Distribuição da ação": "DATA_ACAO",
        "Designação de perícia médica judicial": "DATA_PERICIA",  # Não tem coluna específica
        "Realização da perícia médica": "DATA_REALIZADA", 
        "Laudo pericial": "DATA_LAUDO",
        "Sentença Proferida": "DATA_SENTENCA",
        "Apelação interposta": "DATA_APELACAO",
        "Julgamento": "DATA_JULGAMENTO", 
        "Trânsito em julgado": "DATA_TRANSITO",
        "Cumprimento de Sentença": "DATA_CUMPRIMENTO",
        "Homologação dos Valores": "DATA_HOMOLOGACAO",
        "Instauração do RPV ou Precatório": "DATA_PRECA",
        "Expedição de Ofício":"DATA_OFICIO",
        "Ordem de Pagamento":"DATA_OR_PAGAMENTO",
        "Encerramento com a Liberação do Valores": "DATA_ENCERRAMENTO"
    }

    # Para cada processo na planilha 
    for index, row in df.iterrows(): 
        numero_processo = str(row['NUMERO_PROCESSO']) if 'NUMERO_PROCESSO' in row else f'_{index+1:03d}' 
        print(f"Processando processo {index + 1}/{len(df)}: {numero_processo}") 
        
        try: 
            # Carregar o template Word 
            doc = Document(arquivo_word) 
            
            # Calcular datas previstas
            data_acao = row.get('DATA_ACAO')
            datas_previstas = calcular_datas_previstas_sequencia(data_acao)
            
            # Função para obter valor com fallback para data prevista
            def obter_valor_com_previsto(coluna, placeholder, evento_correspondente=None):
                valor_real = row[coluna] if coluna in row else None
                
                # Se tem valor real, formata e retorna
                if not (pd.isna(valor_real) or valor_real in ('', 'None', None)):
                    if 'DATA' in coluna:
                        return formatar_data(valor_real), False
                    else:
                        return str(valor_real), False
                
                # Se não tem valor real mas tem data prevista correspondente
                if evento_correspondente and evento_correspondente in datas_previstas:
                    return datas_previstas[evento_correspondente], True
                
                # Não tem nem valor real nem previsto
                return '', False

            # Substituir nos parágrafos 
            for paragraph in doc.paragraphs: 
                texto_original = paragraph.text 
                texto_substituido = texto_original 
                
                # Substituições diretas {CHAVE} 
                for coluna, placeholder in mapeamento.items(): 
                    chave = f'{{{placeholder}}}' 
                    if chave in texto_substituido: 
                        # Verificar se há evento correspondente para data prevista
                        evento_correspondente = None
                        for evento, col in mapeamento_eventos_colunas.items():
                            if col == coluna:
                                evento_correspondente = evento
                                break
                                
                        valor, eh_previsto = obter_valor_com_previsto(coluna, placeholder, evento_correspondente)
                        texto_substituido = texto_substituido.replace(chave, valor) 

                # Marcações especiais -> coloca X dentro dos parênteses quando bater 
                novo_texto = aplicar_marcacoes(texto_substituido, row) 
                
                if novo_texto != texto_original: 
                    # Limpa os runs existentes e insere o novo texto com formatação 
                    for r in paragraph.runs: 
                        r.text = '' 
                    new_run = paragraph.add_run(novo_texto) 
                    aplicar_fonte_calibri_light(new_run) 

            # Substituir nas tabelas 
            for table in doc.tables: 
                for row_table in table.rows: 
                    for cell in row_table.cells: 
                        # Processar cada parágrafo dentro da célula 
                        for paragraph in cell.paragraphs: 
                            texto_original = paragraph.text 
                            texto_substituido = texto_original 
                            
                            # Substituições diretas {CHAVE} 
                            for coluna, placeholder in mapeamento.items(): 
                                chave = f'{{{placeholder}}}' 
                                if chave in texto_substituido: 
                                    # Verificar se há evento correspondente para data prevista
                                    evento_correspondente = None
                                    for evento, col in mapeamento_eventos_colunas.items():
                                        if col == coluna:
                                            evento_correspondente = evento
                                            break
                                            
                                    valor, eh_previsto = obter_valor_com_previsto(coluna, placeholder, evento_correspondente)
                                    
                                    if eh_previsto:
                                        # Para datas previstas, criar run com cor vermelha
                                        for r in paragraph.runs:
                                            r.text = ''
                                        new_run = paragraph.add_run(texto_substituido.replace(chave, valor))
                                        aplicar_fonte_calibri_light(new_run, cor_vermelha=True)
                                    else:
                                        texto_substituido = texto_substituido.replace(chave, valor)

                            # Marcações especiais -> coloca X dentro dos parênteses quando bater 
                            texto_substituido = aplicar_marcacoes(texto_substituido, row) 
                            
                            if texto_substituido != texto_original: 
                                # Limpa os runs existentes e insere o novo texto com formatação 
                                for r in paragraph.runs: 
                                    r.text = '' 
                                new_run = paragraph.add_run(texto_substituido) 
                                aplicar_fonte_calibri_light(new_run) 

            # Salvar o documento preenchido 
            nome_arquivo_saida = f'RELATORIO_{limpar_nome_arquivo(numero_processo)}.docx' 
            caminho_saida = os.path.join(diretorio_atual, nome_arquivo_saida) 
            doc.save(caminho_saida) 
            print(f'✓ Relatório gerado: {nome_arquivo_saida}') 
            
        except Exception as e: 
            print(f"✗ Erro ao processar processo {numero_processo}: {e}") 

    print("\nProcessamento concluído!") 

if __name__ == "__main__": 
    preencher_relatorio()