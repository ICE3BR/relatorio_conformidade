
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from datetime import datetime
from dateutil.relativedelta import relativedelta
import os
import re


def preencher_relatorio():
    # Definir caminhos
    diretorio_atual = os.path.dirname(os.path.abspath(__file__))
    arquivo_excel = os.path.join(diretorio_atual, 'Conformidade.xlsx')
    arquivo_modelo_preca = os.path.join(diretorio_atual, 'MODELO RELATORIO.docx')
    arquivo_modelo_rpv = os.path.join(diretorio_atual, 'Conformidade  - RPV.docx')

    # Verificar arquivos
    if not os.path.exists(arquivo_excel):
        print(f"ERRO: Arquivo Excel não encontrado: {arquivo_excel}")
        return
    if not os.path.exists(arquivo_modelo_preca):
        print(f"ERRO: Arquivo Word PRECA não encontrado: {arquivo_modelo_preca}")
        return
    if not os.path.exists(arquivo_modelo_rpv):
        print(f"ERRO: Arquivo Word RPV não encontrado: {arquivo_modelo_rpv}")
        return

    # Ler Excel
    try:
        df = pd.read_excel(arquivo_excel)
        print(f"Planilha carregada com {len(df)} processos")
    except Exception as e:
        print(f"Erro ao ler a planilha Excel: {e}")
        return

    # Placeholders {CHAVE} -> colunas
    mapeamento = {
        'NUMERO_PROCESSO': 'NUMERO_PROCESSO',
        'AUTOR': 'AUTOR',
        'CUMPRIMENTO_SENTENCA': 'CUMPRIMENTO_SENTENCA',
        'SITUACAO_PROCESSO': 'SITUACAO_PROCESSO',
        'DATA_ACAO': 'DATA_ACAO',
        'DATA_PERICIA': 'DATA_PERICIA',
        'DATA_REALIZADA': 'DATA_REALIZADA',
        'DATA_LAUDO': 'DATA_LAUDO',
        'TIPO LAUDO': 'TIPO LAUDO',
        'DATA_SENTENCA': 'DATA_SENTENCA',
        'SENTENCA': 'SENTENCA',
        'DATA_APELACAO': 'DATA_APELACAO',
        'APE': 'APE',
        'DATA_JULGAMENTO': 'DATA_JULGAMENTO',
        'JULGA': 'JULGA',
        'DATA_TRANSITO': 'DATA_TRANSITO',
        'DATA_CUMPRIMENTO': 'DATA_CUMPRIMENTO',
        'DATA_HOMOLOGACAO': 'DATA_HOMOLOGACAO',
        'DATA_PRECA': 'DATA_PRECA',
        'DATA_RPV': 'DATA_RPV',
        'DATA_OFICIO': 'DATA_OFICIO',
        'DATA_OR_PAGAMENTO': 'DATA_OR_PAGAMENTO',
        'DATA_ENCERRAMENTO': 'DATA_ENCERRAMENTO'
    }

    # Sequências diferentes para PRECA e RPV
    EVENTOS_SEQUENCIA_PRECA = [
        ("DATA_ACAO", 0, 0, 0),
        ("DATA_PERICIA", 0, 2, 0),
        ("DATA_REALIZADA", 0, 1, 20),
        ("DATA_LAUDO", 0, 2, 0),
        ("DATA_SENTENCA", 0, 3, 0),
        ("DATA_APELACAO", 0, 1, 15),
        ("DATA_JULGAMENTO", 0, 2, 0),
        ("DATA_TRANSITO", 0, 1, 0),
        ("DATA_CUMPRIMENTO", 0, 1, 1),
        ("DATA_HOMOLOGACAO", 0, 3, 0),
        ("DATA_PRECA", 0, 1, 5),
        ("DATA_OFICIO", 0, 3, 0),
        ("DATA_OR_PAGAMENTO", 0, 1, 0),
        ("DATA_ENCERRAMENTO", 1, 6, 0),
    ]

    EVENTOS_SEQUENCIA_RPV = [
        ("DATA_ACAO", 0, 0, 0),
        ("DATA_PERICIA", 0, 2, 0),
        ("DATA_REALIZADA", 0, 1, 20),
        ("DATA_LAUDO", 0, 2, 0),
        ("DATA_SENTENCA", 0, 3, 0),
        ("DATA_APELACAO", 0, 1, 15),
        ("DATA_JULGAMENTO", 0, 2, 0),
        ("DATA_TRANSITO", 0, 1, 0),
        ("DATA_CUMPRIMENTO", 0, 1, 1),
        ("DATA_HOMOLOGACAO", 0, 3, 0),
        ("DATA_RPV", 0, 1, 5),
        ("DATA_ENCERRAMENTO", 0, 3, 9),
    ]

    # -------- Helpers --------
    def _parse_data(valor):
        """Converte para date (ou None) tentando dd/mm/yyyy e yyyy-mm-dd."""
        if pd.isna(valor) or valor in ('', None, 'None'):
            return None
        try:
            return pd.to_datetime(valor, dayfirst=True, errors='raise').date()
        except Exception:
            try:
                return pd.to_datetime(valor, dayfirst=False, errors='raise').date()
            except Exception:
                return None

    def _fmt_dt(dt):
        return '' if dt is None else dt.strftime('%d/%m/%Y')

    def formatar_data(valor):
        return _fmt_dt(_parse_data(valor))

    def limpar_nome_arquivo(nome):
        for ch in ['<', '>', ':', '"', '/', '\\', '|', '?', '*']:
            nome = nome.replace(ch, '-')
        return nome

    def norm(txt: str) -> str:
        """Normaliza para comparação (lowercase; remove acentos; trata NaN)."""
        if txt is None or (isinstance(txt, float) and pd.isna(txt)):
            return ''
        s = str(txt).strip().lower()
        s = (s.replace('á', 'a').replace('à', 'a').replace('â', 'a').replace('ã', 'a')
               .replace('é', 'e').replace('ê', 'e')
               .replace('í', 'i')
               .replace('ó', 'o').replace('ô', 'o').replace('õ', 'o')
               .replace('ú', 'u')
               .replace('ç', 'c'))
        return s

    def aplicar_marcacoes(texto: str, linha: pd.Series) -> str:
        """Preenche tokens especiais com (X) / ( ). Nunca chama .lower() direto em valores."""
        if not texto:
            return texto

        marcacoes = {
            'LP': '( )', 'LPP': '( )', 'LN': '( )',
            'SENTENCA_A': '( )', 'SENTENCA_I': '( )',
            'APE_A': '( )', 'APE_I': '( )',
            'JULGA_A': '( )', 'JULGA_I': '( )',
        }

        # Laudo
        laudo_bruto = linha.get('TIPO LAUDO', '') or linha.get('LAUDO', '')
        laudo_n = norm(laudo_bruto)
        if 'positivo' in laudo_n:
            marcacoes['LP'] = '(X)'
        elif 'parcial' in laudo_n:
            marcacoes['LPP'] = '(X)'
        elif 'negativo' in laudo_n:
            marcacoes['LN'] = '(X)'

        # Sentença
        sentenca_bruto = linha.get('SENTENÇA', '') or linha.get('SENTENCA', '')
        sentenca_n = norm(sentenca_bruto)
        if 'procedente' in sentenca_n:
            marcacoes['SENTENCA_A'] = '(X)'
        elif 'improcedent e' in sentenca_n:
            marcacoes['SENTENCA_I'] = '(X)'

        # Apelação
        apelacao_bruto = linha.get('APELAÇÃO', '') or linha.get('APELACAO', linha.get('APE', ''))
        apelacao_n = norm(apelacao_bruto)
        if 'autor' in apelacao_n:
            marcacoes['APE_A'] = '(X)'
        elif 'inss' in apelacao_n:
            marcacoes['APE_I'] = '(X)'

        # Julgamento
        julgamento_bruto = linha.get('JULGAMENTO', '') or linha.get('JULGA', '')
        julgamento_n = norm(julgamento_bruto)
        if 'favoravel' in julgamento_n:
            marcacoes['JULGA_A'] = '(X)'
        elif 'desfavoravel' in julgamento_n:
            marcacoes['JULGA_I'] = '(X)'

        for token, repl in marcacoes.items():
            texto = texto.replace(f'({{{token}}})', repl)
        return texto

    def aplicar_fonte_calibri_light(run, cor_vermelha=False):
        run.font.name = 'Calibri Light'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Calibri Light')
        run.font.size = Pt(10.5)
        if cor_vermelha:
            run.font.color.rgb = RGBColor(255, 0, 0)

    def resolver_datas(row, sequencia_eventos):
        """
        Retorna: {'COLUNA': {'valor': 'DD/MM/AAAA'|'', 'prevista': bool}}
        1) Acha a ÚLTIMA data REAL cronologicamente; 2) mantém as reais até ela;
        3) prevê só o que vem depois usando relativedelta(years/months/days).
        """
        datas = {}
        reais = []
        for i, (col, _, _, _) in enumerate(sequencia_eventos):
            dt = _parse_data(row.get(col))
            if dt is not None:
                reais.append((i, col, dt))

        if not reais:
            for col, _, _, _ in sequencia_eventos:
                datas[col] = {'valor': '', 'prevista': False}
            return datas

        idx_anchor, col_anchor, dt_anchor = max(reais, key=lambda t: t[2])

        for i, (col, _, _, _) in enumerate(sequencia_eventos[:idx_anchor + 1]):
            dt = _parse_data(row.get(col))
            datas[col] = {'valor': _fmt_dt(dt) if dt else '', 'prevista': False}

        cursor = dt_anchor
        for i in range(idx_anchor + 1, len(sequencia_eventos)):
            col, anos, meses, dias = sequencia_eventos[i]
            dt_real = _parse_data(row.get(col))
            if dt_real:
                datas[col] = {'valor': _fmt_dt(dt_real), 'prevista': False}
                if dt_real > cursor:
                    cursor = dt_real
            else:
                cursor = cursor + relativedelta(years=anos, months=meses, days=dias)
                datas[col] = {'valor': _fmt_dt(cursor), 'prevista': True}
        return datas

    def determinar_modelos(row):
        """Determina quais modelos usar baseado nas datas existentes"""
        data_rpv = _parse_data(row.get('DATA_RPV'))
        data_preca = _parse_data(row.get('DATA_PRECA'))
        
        modelos = []
        
        # Se tem data RPV, usa modelo RPV
        if data_rpv is not None:
            modelos.append(('RPV', arquivo_modelo_rpv, EVENTOS_SEQUENCIA_RPV))
        
        # Se tem data PRECA, usa modelo PRECA
        if data_preca is not None:
            modelos.append(('PRECA', arquivo_modelo_preca, EVENTOS_SEQUENCIA_PRECA))
        
        # Se não tem nenhum dos dois, usa PRECA como padrão
        if not modelos:
            modelos.append(('PRECA', arquivo_modelo_preca, EVENTOS_SEQUENCIA_PRECA))
            
        return modelos

    def preencher_documento(doc, row, datas_resolvidas):
        """Preenche um documento com os dados da linha"""
        # Parágrafos
        for paragraph in doc.paragraphs:
            texto_original = paragraph.text
            texto_substituido = texto_original
            tem_prevista = False

            for coluna, placeholder in mapeamento.items():
                chave = f'{{{placeholder}}}'
                if chave in texto_substituido:
                    if 'DATA' in coluna:
                        info = datas_resolvidas.get(coluna, {'valor': '', 'prevista': False})
                        texto_substituido = texto_substituido.replace(chave, info['valor'])
                        tem_prevista = tem_prevista or (info['prevista'] and info['valor'] != '')
                    else:
                        valor = row.get(coluna)
                        valor = '' if pd.isna(valor) or valor in ('', None, 'None') else str(valor)
                        texto_substituido = texto_substituido.replace(chave, valor)

            texto_substituido = aplicar_marcacoes(texto_substituido, row)

            if texto_substituido != texto_original:
                for r in paragraph.runs:
                    r.text = ''
                new_run = paragraph.add_run(texto_substituido)
                aplicar_fonte_calibri_light(new_run, cor_vermelha=tem_prevista)

        # Tabelas
        for table in doc.tables:
            for row_table in table.rows:
                for cell in row_table.cells:
                    for paragraph in cell.paragraphs:
                        texto_original = paragraph.text
                        texto_substituido = texto_original
                        tem_prevista = False

                        for coluna, placeholder in mapeamento.items():
                            chave = f'{{{placeholder}}}'
                            if chave in texto_substituido:
                                if 'DATA' in coluna:
                                    info = datas_resolvidas.get(coluna, {'valor': '', 'prevista': False})
                                    texto_substituido = texto_substituido.replace(chave, info['valor'])
                                    tem_prevista = tem_prevista or (info['prevista'] and info['valor'] != '')
                                else:
                                    valor = row.get(coluna)
                                    valor = '' if pd.isna(valor) or valor in ('', None, 'None') else str(valor)
                                    texto_substituido = texto_substituido.replace(chave, valor)

                        texto_substituido = aplicar_marcacoes(texto_substituido, row)

                        if texto_substituido != texto_original:
                            for r in paragraph.runs:
                                r.text = ''
                            new_run = paragraph.add_run(texto_substituido)
                            aplicar_fonte_calibri_light(new_run, cor_vermelha=tem_prevista)

        return doc

    # -------- Processar cada linha --------
    for index, row in df.iterrows():
        numero_processo = str(row['NUMERO_PROCESSO']) if 'NUMERO_PROCESSO' in row else f'_{index+1:03d}'
        print(f"Processando processo {index + 1}/{len(df)}: {numero_processo}")

        try:
            # Determinar quais modelos usar
            modelos_usar = determinar_modelos(row)
            print(f"  Modelos a gerar: {[modelo[0] for modelo in modelos_usar]}")
            
            for tipo_modelo, arquivo_modelo, sequencia in modelos_usar:
                doc = Document(arquivo_modelo)
                datas_resolvidas = resolver_datas(row, sequencia)
                
                # Preencher documento
                doc = preencher_documento(doc, row, datas_resolvidas)

                # Salvar
                nome_arquivo_saida = f'RELATORIO_{limpar_nome_arquivo(numero_processo)}_{tipo_modelo}.docx'
                doc.save(os.path.join(diretorio_atual, nome_arquivo_saida))
                print(f'  ✓ Relatório {tipo_modelo} gerado: {nome_arquivo_saida}')

        except Exception as e:
            print(f"✗ Erro ao processar processo {numero_processo}: {e}")

    print("\nProcessamento concluído!")


if __name__ == "__main__":
    preencher_relatorio()
